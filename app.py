from flask import Flask, render_template, request, send_file
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate_docx', methods=['POST'])
def generate_docx():
    title = request.form['title']
    abstract = request.form['abstract']
    introduction = request.form['introduction']
    objectives = request.form.getlist('objectives[]')
    
    doc = Document()
    title_heading=doc.add_heading()
    title_run=title_heading.add_run(title)
    title_run.bold = True
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_heading("ABSTRACT", level=2)
    para=doc.add_paragraph(abstract)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_heading("INTRODUCTION", level=2)
    para=doc.add_paragraph(introduction)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    
    doc.add_heading('Objectives', level=2)
    for obj in objectives:
        doc.add_paragraph(f"- {obj}")

    doc_file = 'generated_report.docx'
    doc.save(doc_file)

    return send_file(doc_file, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
