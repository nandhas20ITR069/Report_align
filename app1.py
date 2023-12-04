from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import PySimpleGUI as sg

def create_project_report(report_data):
    doc = Document()

    # Title
    title_text = report_data.get('Title', '')
    title = doc.add_heading(title_text, level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Abstract
    abstract_text = report_data.get('Abstract', '')
    doc.add_heading('Abstract', level=2)
    doc.add_paragraph(abstract_text)

    # Introduction
    intro_text = report_data.get('Introduction', '')
    doc.add_heading('Introduction', level=2)
    doc.add_paragraph(intro_text)

    # Objectives
    doc.add_heading('Objectives', level=2)
    objectives_text = report_data.get('Objectives', '').split(',')
    for obj in objectives_text:
        doc.add_paragraph(f"- {obj.strip()}")

    # Methodology (similarly for other sections)
    # ...

    doc.save('User_Generated_Report.docx')

def main():
    sg.theme('LightGrey1')

    layout = [
        [sg.Text('Project Title:'), sg.InputText(key='Title')],
        [sg.Text('Abstract:'), sg.InputText(key='Abstract')],
        [sg.Text('Introduction:'), sg.InputText(key='Introduction')],
        [sg.Text('Objectives (separate by commas):'), sg.InputText(key='Objectives')],
        [sg.Button('Add'), sg.Button('Generate Report')]
    ]

    window = sg.Window('Project Report Generator', layout)

    report_data = {}

    while True:
        event, values = window.read()

        if event == sg.WINDOW_CLOSED:
            break

        if event == 'Add':
            section = sg.popup_get_text('Enter Section Name:')
            if section:
                text = sg.popup_get_text(f'Enter {section}:')
                if text:
                    report_data[section] = text

        if event == 'Generate Report':
            create_project_report(report_data)
            sg.popup('Report Generated!', 'User_Generated_Report.docx')

    window.close()

if __name__ == "__main__":
    main()
